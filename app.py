from flask import Flask, request, jsonify, render_template
from outline import generate_response, generate_diagram

app = Flask(__name__)

# Add to app.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from flask import send_file
import markdown
from bs4 import BeautifulSoup

@app.route('/download', methods=['POST'])
def download_word():
    data = request.json
    
    if not data or 'content' not in data:
        return jsonify({'error': 'No content provided'}), 400
        
    markdown_content = data['content']
    print(markdown_content)
    
    try:
        # Convert markdown to HTML for better parsing
        html_content = markdown.markdown(markdown_content)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create a new Word document
        doc = Document()
        
        # Set up document styles
        def create_heading_style(doc, level):
            style_name = f'CustomHeading{level}'
            if style_name not in doc.styles:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = doc.styles[f'Heading {level}']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(20 - (level * 2))  # Decrease size for deeper headings
                font.bold = True
                paragraph_format = style.paragraph_format
                paragraph_format.space_before = Pt(12)
                paragraph_format.space_after = Pt(6)
                paragraph_format.line_spacing = 1.15
            return style_name

        # Set up normal paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15
        
        # Set up list styles
        doc.styles['List Bullet'].paragraph_format.left_indent = Inches(0.5)
        doc.styles['List Number'].paragraph_format.left_indent = Inches(0.5)
        
        # Process HTML elements and add them to the Word document
        current_list = None
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                heading_style = create_heading_style(doc, level)
                paragraph = doc.add_paragraph(style=heading_style)
                paragraph.add_run(element.get_text().strip())
                current_list = None
                
            elif element.name == 'p':
                if not element.get_text().strip():  # Skip empty paragraphs
                    continue
                paragraph = doc.add_paragraph(style='Normal')
                paragraph.add_run(element.get_text().strip())
                current_list = None
                
            elif element.name == 'ul':
                current_list = 'bullet'
                continue
                
            elif element.name == 'ol':
                current_list = 'number'
                continue
                
            elif element.name == 'li' and current_list:
                list_style = 'List Bullet' if current_list == 'bullet' else 'List Number'
                paragraph = doc.add_paragraph(style=list_style)
                paragraph.add_run(element.get_text().strip())
        
        # Add page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='outline.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/outline', methods=['POST'])
def outline():
    data = request.json
    
    # Check if outline exists in the request data
    if not data or 'outline' not in data:
        return jsonify({'error': 'No outline text provided'}), 400
        
    outline_text = data['outline']
    
    # Check if the outline text is empty
    if not outline_text or not isinstance(outline_text, str):
        return jsonify({'error': 'Cannot generate outline without text'}), 400
    
    try:
        # Process the outline text
        result = generate_response(outline_text)
        return jsonify({'result': result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
# Add to app.py

@app.route('/generate-diagram', methods=['POST'])
def generate_diagram_route():  # Renamed to avoid conflict with imported function
    data = request.json
    
    if not data or 'type' not in data or 'outline' not in data:
        return jsonify({'error': 'No diagram type provided or outline was not generated'}), 400
    
    diagram_type = data['type']
    instructions = data.get('instructions', '')
    outline = data['outline']

    diagram_input = f"Diagram Type: {diagram_type}\n\n\n\nDiagram Instructions:\n{instructions}\n\n\n\nOutline:\n{outline}"
    
    try:
        # Process the diagram text
        result = generate_diagram(diagram_input)  # Calling the imported function
        return jsonify({'result': result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)